import asyncio
import json
import os
import re
import tempfile
from pathlib import Path
from typing import List, Optional

import pdfplumber
from docx import Document
from dotenv import load_dotenv
from emergentintegrations.llm.chat import LlmChat, UserMessage
from pydantic import BaseModel, Field, ValidationError
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / ".env")


class SectionFeedback(BaseModel):
    section: str
    score: int = Field(ge=0, le=10)
    strengths: List[str] = Field(default_factory=list)
    issues: List[str] = Field(default_factory=list)
    recommendations: List[str] = Field(default_factory=list)


class RewriteExample(BaseModel):
    original: Optional[str] = None
    improved: str
    reason: str


class CVReview(BaseModel):
    overall_score: int = Field(ge=0, le=10)
    summary: str
    section_scores: List[SectionFeedback]
    top_strengths: List[str]
    priority_improvements: List[str]
    rewrite_examples: List[RewriteExample] = Field(default_factory=list)
    privacy_notes: List[str] = Field(default_factory=list)
    finnish_market_resources: List[str] = Field(default_factory=list)


SAMPLE_CV = """
Maya Lehtinen
Helsinki, Finland | maya.lehtinen@example.com | +358 40 123 4567 | linkedin.com/in/mayalehtinen

PROFILE
Customer support specialist with four years of experience in SaaS and retail. Interested in HR coordinator roles.

EXPERIENCE
Customer Support Specialist, Nordic SaaS Oy, Helsinki | 2021-2025
- Helped customers by email and chat.
- Improved help center articles and trained two new team members.
- Worked with product team to report recurring bugs.

Sales Assistant, K-Market, Espoo | 2019-2021
- Served customers and handled cashier tasks.
- Managed shelf inventory during busy periods.

EDUCATION
Bachelor of Business Administration, Haaga-Helia University of Applied Sciences | 2018-2021

SKILLS
CRM systems, Zendesk, Excel, onboarding, communication, teamwork

LANGUAGES
English fluent, Finnish B1, Swedish basic

HOBBIES
Floorball, volunteering at local events
""".strip()


def create_sample_pdf(path: Path, text: str) -> None:
    c = canvas.Canvas(str(path), pagesize=A4)
    _, height = A4
    y = height - 50
    for line in text.splitlines():
        if y < 50:
            c.showPage()
            y = height - 50
        c.drawString(45, y, line[:110])
        y -= 15
    c.save()


def create_sample_docx(path: Path, text: str) -> None:
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.save(str(path))


def extract_pdf(path: Path) -> str:
    try:
        with pdfplumber.open(str(path)) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        if text.strip():
            return text.strip()
    except Exception as exc:
        print(f"pdfplumber failed, falling back to pypdf: {exc}")

    reader = PdfReader(str(path))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    return text.strip()


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()


def extract_text(path: Optional[Path], fallback_text: str = "") -> str:
    if not path:
        return fallback_text.strip()
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    raise ValueError(f"Unsupported file type: {suffix}")


def parse_json_response(raw: str) -> dict:
    cleaned = raw.strip()
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start == -1 or end == -1:
        raise ValueError(f"No JSON object found in response: {raw[:300]}")
    return json.loads(cleaned[start : end + 1])


def build_prompt(cv_text: str, job_title: str, industry: str) -> str:
    trimmed = cv_text[:12000]
    return f"""
Review this CV for a person seeking employment in Finland.
Target job title: {job_title or "Not specified"}
Target industry: {industry or "Not specified"}

Return ONLY valid JSON matching exactly this shape:
{{
  "overall_score": 0,
  "summary": "short constructive summary",
  "section_scores": [
    {{"section":"Personal Information and Contact Details","score":0,"strengths":[],"issues":[],"recommendations":[]}},
    {{"section":"Format and Structure","score":0,"strengths":[],"issues":[],"recommendations":[]}},
    {{"section":"Professional Experience","score":0,"strengths":[],"issues":[],"recommendations":[]}},
    {{"section":"Education","score":0,"strengths":[],"issues":[],"recommendations":[]}},
    {{"section":"Skills and Competencies","score":0,"strengths":[],"issues":[],"recommendations":[]}},
    {{"section":"Language Proficiency","score":0,"strengths":[],"issues":[],"recommendations":[]}},
    {{"section":"Hobbies and Interests","score":0,"strengths":[],"issues":[],"recommendations":[]}},
    {{"section":"Cultural Fit and Adaptability","score":0,"strengths":[],"issues":[],"recommendations":[]}},
    {{"section":"General Recommendations","score":0,"strengths":[],"issues":[],"recommendations":[]}}
  ],
  "top_strengths": [],
  "priority_improvements": [],
  "rewrite_examples": [{{"original":"optional copied weak bullet","improved":"improved STAR-style bullet","reason":"why this is better"}}],
  "privacy_notes": [],
  "finnish_market_resources": []
}}

Rules:
- Scores must be integers 0-10.
- Every section must include at least one strength or issue and at least two actionable recommendations.
- Advice must be specific to Finland: privacy expectations, concise 1-2 page norms, Finnish/Swedish language levels using CEFR A1-C2, equality/teamwork/direct communication, local education references if relevant.
- Do not invent personal facts. If information is missing, say it is missing.
- Include practical resources such as TE-palvelut/Job Market Finland, LinkedIn Finland, Duunitori, Oikotie Työpaikat where relevant.

CV text:
---
{trimmed}
---
""".strip()


async def run_review(cv_text: str, job_title: str = "HR Coordinator", industry: str = "Human Resources") -> CVReview:
    api_key = os.environ.get("EMERGENT_LLM_KEY")
    if not api_key:
        raise RuntimeError("EMERGENT_LLM_KEY missing")
    chat = LlmChat(
        api_key=api_key,
        session_id="cv-reviewer-fi-poc",
        system_message="You are an expert Finnish job-market CV reviewer. You produce strict JSON only.",
    ).with_model("openai", "gpt-5.2")
    raw = await chat.send_message(UserMessage(text=build_prompt(cv_text, job_title, industry)))
    parsed = parse_json_response(raw)
    return CVReview(**parsed)


async def main() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        pdf_path = tmp_dir / "sample_cv.pdf"
        docx_path = tmp_dir / "sample_cv.docx"
        create_sample_pdf(pdf_path, SAMPLE_CV)
        create_sample_docx(docx_path, SAMPLE_CV)

        cases = [
            ("pasted_text", None, SAMPLE_CV),
            ("pdf_upload", pdf_path, ""),
            ("docx_upload", docx_path, ""),
        ]

        for name, path, fallback in cases:
            extracted = extract_text(path, fallback)
            assert len(extracted) > 250, f"{name}: extracted text too short"
            assert "Maya Lehtinen" in extracted, f"{name}: missing expected name after extraction"
            print(f"{name}: extraction OK ({len(extracted)} chars)")

            review = await run_review(extracted)
            assert len(review.section_scores) >= 9, f"{name}: missing sections"
            assert review.overall_score >= 0, f"{name}: invalid overall score"
            print(f"{name}: AI review OK | overall={review.overall_score} | sections={len(review.section_scores)}")

    print("CORE_POC_SUCCESS")


if __name__ == "__main__":
    try:
        asyncio.run(main())
    except (ValidationError, Exception) as exc:
        print(f"CORE_POC_FAILED: {exc}")
        raise
